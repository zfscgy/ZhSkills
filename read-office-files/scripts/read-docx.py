"""
read-docx.py  —  读取 .docx 文件，输出结构化文本
用法:
    python read-docx.py <file.docx> [--json] [--tables] [--meta]

选项:
    --json     以 JSON 格式输出（便于程序处理）
    --tables   同时提取表格内容
    --meta     同时输出文档元数据（标题、作者、日期等）

依赖: pip install python-docx
"""

import sys
import json
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


HEADING_PREFIX = {1: "# ", 2: "## ", 3: "### ", 4: "#### "}


def get_paragraph_style_level(para) -> int | None:
    """返回标题级别（1-4），非标题返回 None。"""
    name = para.style.name
    for lvl in (1, 2, 3, 4):
        if name == f"Heading {lvl}":
            return lvl
    return None


def extract_paragraphs(doc) -> list[dict]:
    """提取所有段落，含层级信息。"""
    result = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        level = get_paragraph_style_level(para)
        result.append({
            "type": "heading" if level else "paragraph",
            "level": level,
            "text": text,
        })
    return result


def extract_tables(doc) -> list[list[list[str]]]:
    """提取所有表格，返回 [table][row][col] 结构。"""
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


def extract_meta(doc) -> dict:
    """提取文档核心属性（元数据）。"""
    cp = doc.core_properties
    return {
        "title":    cp.title or "",
        "author":   cp.author or "",
        "subject":  cp.subject or "",
        "keywords": cp.keywords or "",
        "created":  str(cp.created) if cp.created else "",
        "modified": str(cp.modified) if cp.modified else "",
        "revision": cp.revision,
    }


def to_markdown(paragraphs: list[dict], tables: list | None = None) -> str:
    """将提取结果渲染为 Markdown 文本。"""
    lines = []
    for item in paragraphs:
        if item["type"] == "heading":
            prefix = HEADING_PREFIX.get(item["level"], "#### ")
            lines.append(f"{prefix}{item['text']}")
        else:
            lines.append(item["text"])
        lines.append("")  # 段落间空行

    if tables:
        lines.append("---")
        lines.append("## 表格内容")
        for i, table in enumerate(tables, 1):
            lines.append(f"\n### 表格 {i}")
            if not table:
                continue
            # Markdown 表格
            header = table[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in table[1:]:
                lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="读取 .docx 文件内容")
    parser.add_argument("file", help=".docx 文件路径")
    parser.add_argument("--json",   action="store_true", help="JSON 格式输出")
    parser.add_argument("--tables", action="store_true", help="提取表格")
    parser.add_argument("--meta",   action="store_true", help="输出元数据")
    args = parser.parse_args()

    path = Path(args.file)
    if not path.exists():
        print(f"[错误] 文件不存在：{path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(path))
    paragraphs = extract_paragraphs(doc)
    tables = extract_tables(doc) if args.tables else None
    meta = extract_meta(doc) if args.meta else None

    if args.json:
        output = {"paragraphs": paragraphs}
        if tables is not None:
            output["tables"] = tables
        if meta is not None:
            output["meta"] = meta
        print(json.dumps(output, ensure_ascii=False, indent=2))
    else:
        if meta:
            print("=== 文档元数据 ===")
            for k, v in meta.items():
                if v:
                    print(f"{k}: {v}")
            print()
        print(to_markdown(paragraphs, tables))


if __name__ == "__main__":
    main()
