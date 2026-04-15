"""
docx-example.py
完整示例：使用 python-docx 生成一份带有多级标题、正文、表格、
列表、页眉页脚的标准 Word 文档。

依赖：pip install python-docx

──────────────────────────────────────────────────────────────────
中文字体说明（重要）
──────────────────────────────────────────────────────────────────
推荐方案（最简单）：
    在 Word 里把 Heading 1/2/3、Normal、List Bullet 等样式的字体
    手动设置好（如宋体/微软雅黑），保存为 template.docx，然后：
        doc = Document("template.docx")
    加载后样式字体直接继承，无需任何 XML 操作。

本脚本使用的方案（无模板）：
    python-docx 默认模板的 Heading/List 样式含主题字体属性
    （w:asciiTheme="majorHAnsi"），优先级高于 font.name，
    必须通过 XML 显式清除后中文字体才能生效。
    init_styles(doc) 函数完成这项工作。
──────────────────────────────────────────────────────────────────
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 全局字体常量 ────────────────────────────────────────────────
FONT      = "宋体"    # 中英文统一字体（宋体兼容中英数字）
BODY_SIZE = 12
H1_SIZE   = 16
H2_SIZE   = 14
H3_SIZE   = 12

COLOR_TITLE  = RGBColor(0x1F, 0x49, 0x7D)   # 深蓝
COLOR_H2     = RGBColor(0x37, 0x64, 0x92)   # 中蓝
COLOR_BODY   = RGBColor(0x00, 0x00, 0x00)   # 黑色
COLOR_ACCENT = RGBColor(0xC0, 0x00, 0x00)   # 强调红


# ── 工具函数 ────────────────────────────────────────────────────

def set_run_font(run, font=FONT, size=BODY_SIZE,
                 bold=False, italic=False, color: RGBColor | None = None):
    """设置 run 字体，同时清除主题字体覆盖，确保中文正确显示。"""
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rFonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn("w:ascii"),    font)
    rFonts.set(qn("w:hAnsi"),    font)
    rFonts.set(qn("w:eastAsia"), font)
    # 主题字体优先级高于显式字体名，必须删除否则设置无效
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme"):
        rFonts.attrib.pop(qn(attr), None)


def _apply_style_font(style, font=FONT, size=None, bold=None, color=None,
                      space_before=None, space_after=None, keep_with_next=False):
    """修改段落样式的字体（同时清除主题字体），可选段落间距。"""
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"),    font)
    rFonts.set(qn("w:hAnsi"),    font)
    rFonts.set(qn("w:eastAsia"), font)
    # Heading 样式默认含 w:asciiTheme="majorHAnsi" 等属性，优先级高于显式字体，必须删除
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme"):
        rFonts.attrib.pop(qn(attr), None)
    pf = style.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after  = Pt(space_after)
    if keep_with_next:
        pf.keep_with_next = True


def init_styles(doc):
    """在文档开头统一设置所有常用样式的字体，之后无需逐 run 处理。"""
    _apply_style_font(doc.styles["Heading 1"], size=H1_SIZE, bold=True,
                      color=COLOR_TITLE, space_before=18, space_after=9,
                      keep_with_next=True)
    _apply_style_font(doc.styles["Heading 2"], size=H2_SIZE, bold=True,
                      color=COLOR_H2, space_before=12, space_after=6,
                      keep_with_next=True)
    _apply_style_font(doc.styles["Heading 3"], size=H3_SIZE, bold=True,
                      space_before=9, space_after=4, keep_with_next=True)
    for list_style in ("List Bullet", "List Bullet 2",
                       "List Number", "List Number 2"):
        if list_style in doc.styles:
            _apply_style_font(doc.styles[list_style], size=BODY_SIZE)


def add_body_paragraph(doc, text: str, first_line_indent=True):
    """添加标准正文段落（两端对齐、首行缩进、1.5 倍行距）。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment      = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after    = Pt(6)
    pf.line_spacing   = 1.5
    if first_line_indent:
        pf.first_line_indent = Pt(BODY_SIZE * 2)
    run = p.add_run(text)
    set_run_font(run, color=COLOR_BODY)
    return p


def shade_cell(cell, fill_hex: str = "1F497D"):
    """为表格单元格设置背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_page_number(paragraph):
    """在段落末尾插入"第 X 页  共 Y 页"域代码。"""
    def _field(instr: str):
        r = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        inst = OxmlElement("w:instrText")
        inst.text = f" {instr} "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in (begin, inst, end):
            r._r.append(el)

    paragraph.add_run("第 ")
    _field("PAGE")
    paragraph.add_run(" 页  共 ")
    _field("NUMPAGES")
    paragraph.add_run(" 页")


# ── 主函数 ──────────────────────────────────────────────────────

def build_document() -> Document:
    doc = Document()

    # 页面设置：A4，标准页边距
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.17)
    sec.right_margin  = Cm(3.17)

    # 统一初始化所有样式字体（标题 + 列表），必须在添加内容前调用
    init_styles(doc)

    # ── 封面标题 ────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(72)
    title_p.paragraph_format.space_after  = Pt(18)
    title_run = title_p.add_run("人工智能技术应用报告")
    set_run_font(title_run, size=22, bold=True, color=COLOR_TITLE)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(60)
    sub_run = subtitle_p.add_run("2026年度综合研究")
    set_run_font(sub_run, size=14, color=COLOR_H2)

    doc.add_page_break()

    # ── 第一章 ──────────────────────────────────────────────────
    doc.add_heading("第一章  研究背景与意义", level=1)

    doc.add_heading("1.1  研究背景", level=2)
    add_body_paragraph(
        doc,
        "近年来，人工智能技术迎来了新一轮的爆发式增长。以大型语言模型为代表的"
        "生成式AI技术，正在深刻改变各行各业的生产方式与组织形态。"
    )
    add_body_paragraph(
        doc,
        "本报告从技术演进、产业应用及政策环境三个维度，对当前AI发展现状进行"
        "系统性梳理，并对未来趋势提出展望。"
    )

    doc.add_heading("1.2  研究意义", level=2)
    doc.add_heading("1.2.1  理论价值", level=3)
    add_body_paragraph(
        doc,
        "本研究有助于构建AI技术应用的分析框架，为后续学术研究提供方法论参考。"
    )
    doc.add_heading("1.2.2  实践价值", level=3)
    add_body_paragraph(
        doc,
        "研究结论可直接指导企业的AI转型策略，降低技术应用中的不确定性。"
    )

    # ── 第二章 ──────────────────────────────────────────────────
    doc.add_heading("第二章  核心技术概述", level=1)

    doc.add_heading("2.1  主要技术方向", level=2)
    add_body_paragraph(doc, "当前人工智能的核心技术方向主要包括以下几类：",
                        first_line_indent=True)

    for item in ["大型语言模型（LLM）", "多模态感知与生成", "具身智能与机器人",
                 "AI Agent 自主决策"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.2  技术成熟度对比", level=2)
    add_body_paragraph(doc, "下表列举了各主要技术方向的当前成熟度及典型应用场景：")

    # 表格
    headers = ["技术方向", "成熟度", "典型应用", "代表厂商"]
    rows = [
        ["大型语言模型", "★★★★★", "对话、写作、代码生成", "OpenAI / Anthropic"],
        ["多模态生成",   "★★★★☆", "图像生成、视频合成",   "Stability / Google"],
        ["具身智能",     "★★★☆☆", "工业机器人、家庭助手", "Figure / Boston Dynamics"],
        ["AI Agent",    "★★★☆☆", "自动化工作流、RPA",    "AutoGPT / Zapier"],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 表头
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        shade_cell(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 数据行
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            set_run_font(run, size=11)

    # ── 第三章 ──────────────────────────────────────────────────
    doc.add_heading("第三章  产业应用案例", level=1)

    doc.add_heading("3.1  金融行业", level=2)
    add_body_paragraph(
        doc,
        "金融行业已率先将大模型应用于智能客服、风控审核及研报生成等场景。"
        "据统计，头部银行通过引入AI辅助，客服人力成本降低约 35%。"
    )

    doc.add_heading("3.2  医疗健康", level=2)
    add_body_paragraph(
        doc,
        "AI辅助诊断在影像识别领域已达到或超越专科医生水平，在某些癌症早筛"
        "场景中准确率超过 92%。"
    )

    # 有序列表示例
    doc.add_heading("3.3  实施步骤参考", level=2)
    for step in ["需求调研与场景选定", "数据采集与清洗",
                 "模型选型与微调", "系统集成与测试", "上线监控与迭代"]:
        doc.add_paragraph(step, style="List Number")

    # ── 页眉页脚 ────────────────────────────────────────────────
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "人工智能技术应用报告"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(hp.runs[0] if hp.runs else hp.add_run(), size=10,
                 color=RGBColor(0x80, 0x80, 0x80))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)
    for run in fp.runs:
        set_run_font(run, size=10, color=RGBColor(0x80, 0x80, 0x80))

    return doc


if __name__ == "__main__":
    output_path = "output.docx"
    doc = build_document()
    doc.save(output_path)
    print(f"文档已保存至 {output_path}")
